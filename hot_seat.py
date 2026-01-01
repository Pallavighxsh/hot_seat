#!/usr/bin/env python3

import os
import time
import requests
from pathlib import Path
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document   # ← added for .docx export
from docx.shared import Pt   # (for formatting, minimal)
# reportlab removed from use but keeping all else untouched

from llama_cpp import Llama

# ======================================================
# CONFIG (DO NOT CHANGE LLM SETTINGS)
# ======================================================
ROOT = Path(__file__).parent.resolve()

load_dotenv(ROOT / ".env")

SERP_API_KEY = os.getenv("SERP_API_KEY")
if not SERP_API_KEY:
    raise RuntimeError("SERP_API_KEY not found in .env")

MODEL_PATH = ROOT / "Phi-3-mini-4k-instruct-q4.gguf"
if not MODEL_PATH.exists():
    raise RuntimeError("Phi-3 model file not found")

TOPICS_FILE = ROOT / "topics.txt"
OUTPUT_DIR = ROOT / "output" / "mcqs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

HEADERS = {"User-Agent": "Mozilla/5.0 (Local MCQ Generator)"}

# ======================================================
# LOAD LLM  (MATCHES YOUR WORKING APP)
# ======================================================
print("🔧 Loading Phi-3 model…")

llm = Llama(
    model_path=str(MODEL_PATH),
    n_ctx=2048,
    n_gpu_layers=20,
    n_threads=os.cpu_count(),
    verbose=False,
)

print("✅ Model loaded")

# ======================================================
# SERP SEARCH
# ======================================================
def serp_search(topic, n=3):
    r = requests.get(
        "https://serpapi.com/search",
        params={
            "engine": "google",
            "q": f"{topic} article",
            "num": n,
            "api_key": SERP_API_KEY,
        },
        timeout=20,
    )
    r.raise_for_status()

    return [
        o["link"]
        for o in r.json().get("organic_results", [])
        if "link" in o
    ][:n]

# ======================================================
# SCRAPE ARTICLES
# ======================================================
def scrape_article(url):
    try:
        r = requests.get(url, headers=HEADERS, timeout=15)
        r.raise_for_status()
    except:
        return ""

    soup = BeautifulSoup(r.text, "html.parser")

    selectors = [
        "article", ".entry-content", ".post-content",
        ".content", ".main-content"
    ]

    for sel in selectors:
        block = soup.select_one(sel)
        if block:
            text = " ".join(
                p.get_text(" ", strip=True)
                for p in block.find_all("p")
            )
            if len(text) > 500:
                return text[:6000]

    paras = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
    return " ".join(paras)[:5000]

# ======================================================
# MCQ GENERATION
# ======================================================
def generate_mcqs(topic, context):
    prompt = f"""
You are an expert academic question-setter.

Create exactly 50 multiple-choice questions for the topic:
{topic}

Rules:
- 4 options (A–D)
- Only one correct answer
- Mix difficulty levels
- Do not repeat questions
- Base questions strictly on the content
- Include an ANSWER KEY at the end

Format EXACTLY like this:

Q1. Question text
A) ...
B) ...
C) ...
D) ...

ANSWER KEY:
1. A
2. B
...

CONTENT:
{context[:3500]}
"""

    result = llm(
        prompt,
        max_tokens=1800,
        temperature=0.35,
    )

    return result["choices"][0]["text"].strip()

# ======================================================
# WORD DOCX SAVE  ⇦ (only changed part)
# ======================================================
def save_docx(topic, text):
    filename = topic.replace(" ", "_").replace("&", "and") + ".docx"
    path = OUTPUT_DIR / filename

    doc = Document()
    title = doc.add_heading(f"{topic} – MCQ Question Bank", level=1)

    for line in text.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)

    doc.save(path)
    print(f"📝 Saved → {path.name}")

# ======================================================
# MAIN
# ======================================================
def main():
    if not TOPICS_FILE.exists():
        raise RuntimeError("topics.txt not found")

    with open(TOPICS_FILE) as f:
        topics = [t.strip() for t in f if t.strip()]

    for topic in topics:
        print(f"\n📚 Processing: {topic}")

        try:
            links = serp_search(topic)
        except Exception as e:
            print("❌ SERP error:", e)
            continue

        context = ""
        for link in links:
            context += scrape_article(link) + "\n"
            time.sleep(1)

        if len(context) < 800:
            print("⚠️ Insufficient content — skipped")
            continue

        mcqs = generate_mcqs(topic, context)
        save_docx(topic, mcqs)   # ← PDF replaced with DOCX

    print("\n🎉 All MCQs generated successfully")

# ======================================================
# ENTRY
# ======================================================
if __name__ == "__main__":
    main()
